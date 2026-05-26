from __future__ import annotations

from collections import OrderedDict
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "outputs" / "report" / "final_report_guideline_aligned_submission_draft.docx"

SUMMARY_CSV = ROOT / "outputs" / "tables" / "channel_analysis_summary.csv"
FRAME_CSV = ROOT / "outputs" / "tables" / "channel_frame_distribution.csv"
REACTION_CSV = ROOT / "outputs" / "tables" / "channel_audience_reaction_distribution.csv"
IDEOLOGY_CSV = ROOT / "outputs" / "tables" / "channel_ideology_estimates.csv"
TOPIC_CSV = ROOT / "outputs" / "tables" / "topic_summary.csv"
TOPIC_SCRIPT_CSV = ROOT / "outputs" / "tables" / "topic_summary_script.csv"
TEXT_INPUT_CSV = ROOT / "data" / "processed" / "text_analysis_inputs_stage2.csv"


CHANNEL_TYPE_MAP: dict[str, str] = {
    "YTN": "보도·경제전문 채널 계열",
    "연합뉴스TV": "보도·경제전문 채널 계열",
    "한국경제TV": "보도·경제전문 채널 계열",
    "매일경제TV": "보도·경제전문 채널 계열",
    "SBS Biz 뉴스": "보도·경제전문 채널 계열",
    "KBS News": "공영·지상파 방송사 계열",
    "MBCNEWS": "공영·지상파 방송사 계열",
    "SBS 뉴스": "공영·지상파 방송사 계열",
    "JTBC News": "종합편성채널 계열",
    "채널A News": "종합편성채널 계열",
    "뉴스TVCHOSUN": "종합편성채널 계열",
    "MBN News": "종합편성채널 계열",
}


TOPIC_LABEL_MAP: dict[str, str] = {
    "topic_00": "이란·이스라엘 군사 충돌",
    "topic_01": "진행 멘트·일반 설명",
    "topic_02": "트럼프와 미국 대응",
    "topic_03": "호르무즈 해협과 원유 수송",
    "topic_04": "전쟁의 시장·투자 충격",
    "topic_05": "채널 안내·부가 문구",
    "topic_06": "시사 프로그램·플랫폼 문구",
    "topic_07": "최고지도자와 내부 권력",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_default_font(document: Document) -> None:
    for style_name in ["Normal", "Title", "Subtitle", "List Bullet"]:
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "Malgun Gothic"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            style.font.size = Pt(10.5)


def set_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def add_title(document: Document, title: str, dept_lines: Iterable[str]) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(18)

    for line in dept_lines:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(10.5)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(14 if level == 1 else 11.5)


def add_paragraph(document: Document, text: str, *, indent: bool = False) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10.5)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(10.5)


def add_table(document: Document, headers: list[str], rows: list[list[str]], title: str) -> None:
    add_paragraph(document, title)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "DCE6F1")
        for p in hdr[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = "Malgun Gothic"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
                r.font.size = Pt(9.5)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for p in cells[idx].paragraphs:
                for r in p.runs:
                    r.font.name = "Malgun Gothic"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
                    r.font.size = Pt(9.5)

    document.add_paragraph()


def load_data() -> dict[str, object]:
    summary = pd.read_csv(SUMMARY_CSV)
    frame = pd.read_csv(FRAME_CSV)
    reaction = pd.read_csv(REACTION_CSV)
    ideology = pd.read_csv(IDEOLOGY_CSV)
    topic = pd.read_csv(TOPIC_CSV)
    topic_script = pd.read_csv(TOPIC_SCRIPT_CSV)
    text_inputs = pd.read_csv(TEXT_INPUT_CSV, low_memory=False)

    for df in [summary, frame, reaction, ideology]:
        if "channel_name" in df.columns:
            df["channel_name"] = df["channel_name"].astype(str).str.strip()

    summary["channel_type"] = summary["channel_name"].map(CHANNEL_TYPE_MAP)

    published = pd.to_datetime(text_inputs["published_at"], errors="coerce")
    period_start = published.min().strftime("%Y년 %m월 %d일")
    period_end = published.max().strftime("%Y년 %m월 %d일")

    type_summary = (
        summary.groupby("channel_type", as_index=False)
        .agg(
            channel_count=("channel_name", "count"),
            video_count=("video_count", "sum"),
            comment_count=("audience_comment_count", "sum"),
        )
        .sort_values("video_count", ascending=False)
    )

    frame_totals = frame.groupby("primary_frame", as_index=False)["video_count"].sum()
    frame_totals["share"] = frame_totals["video_count"] / frame_totals["video_count"].sum()
    frame_totals = frame_totals.sort_values("video_count", ascending=False)

    reaction_totals = reaction.groupby("primary_reaction", as_index=False)["comment_count"].sum()
    reaction_totals["share"] = reaction_totals["comment_count"] / reaction_totals["comment_count"].sum()
    reaction_totals = reaction_totals.sort_values("comment_count", ascending=False)

    ideology_sorted = ideology.sort_values("ideology_relative_score", ascending=False)

    topic["topic_name"] = topic["topic_label"].map(TOPIC_LABEL_MAP).fillna(topic["topic_label"])
    topic_script["topic_name"] = topic_script["topic_label"].map(TOPIC_LABEL_MAP).fillna(topic_script["topic_label"])

    return {
        "summary": summary,
        "frame": frame,
        "reaction": reaction,
        "ideology": ideology,
        "topic": topic,
        "topic_script": topic_script,
        "period_start": period_start,
        "period_end": period_end,
        "type_summary": type_summary,
        "frame_totals": frame_totals,
        "reaction_totals": reaction_totals,
        "ideology_sorted": ideology_sorted,
        "total_videos": int(summary["video_count"].sum()),
        "total_comments": int(summary["audience_comment_count"].sum()),
    }


def frame_table_rows() -> list[list[str]]:
    return [
        [
            "안보·군사 프레임",
            "군사 공격, 보복, 확전, 방어, 억지력, 안보 위협 등 물리적 충돌과 군사적 긴장을 중심으로 사건을 해석하는 프레임",
            "Semetko & Valkenburg(2000)의 갈등 프레임을 이란 전쟁 맥락에 맞게 재구성",
        ],
        [
            "국제정치·외교 프레임",
            "국가 책임, 외교 협상, 국제법, 국제기구 대응, 동맹 관계 등 외교적·정치적 관계를 중심으로 사건을 다루는 프레임",
            "Semetko & Valkenburg(2000)의 책임귀인 프레임 및 de Vreese(2005)의 issue-specific frame 논의",
        ],
        [
            "경제·에너지 프레임",
            "국제유가, LNG 공급, 물가, 공급망, 해상 물류, 산업 영향 등 전쟁의 경제적 파급효과를 강조하는 프레임",
            "Semetko & Valkenburg(2000)의 경제적 결과 프레임",
        ],
        [
            "투자·시장 프레임",
            "주가, 증시, 환율, 업종 반응, 투자 심리 등 금융시장 차원의 파급효과를 강조하는 프레임",
            "경제·에너지 프레임의 하위 시장 반응을 분리해 재구성",
        ],
        [
            "인도주의·민간피해 프레임",
            "민간인 사망·부상, 난민, 생활 기반 파괴, 의료시설 피해 등 인간적 피해를 중심으로 사건을 조명하는 프레임",
            "Semetko & Valkenburg(2000)의 인간흥미 프레임을 사안 특화",
        ],
        [
            "기타/혼합 프레임",
            "하나의 프레임으로 명확히 분류하기 어렵거나 둘 이상의 프레임이 동시에 강하게 나타나는 경우",
            "본 연구의 데이터 특성상 추가한 보조 범주",
        ],
    ]


def build_document() -> Document:
    data = load_data()
    doc = Document()
    set_default_font(doc)
    set_page(doc)

    add_title(
        doc,
        "전쟁은 어떻게 재구성되는가: 이란 전쟁 관련 한국 유튜브 뉴스 채널의 보도 프레임과 수용자 반응 분석",
        [
            "인천대학교 사회과학대학",
            "창의인재개발학과",
            "주기쁨",
        ],
    )

    add_heading(doc, "I. 서론", 1)
    add_paragraph(
        doc,
        "중동 지역의 군사 충돌은 단순한 국제 뉴스에 그치지 않고 에너지 공급, 해상 물류, 물가, 외교·안보 인식 등과 연결되며 국내 뉴스 소비에도 빠르게 영향을 미친다. 특히 유튜브는 영상, 제목, 설명, 댓글이 한 공간에 결합된 형태로 뉴스가 유통되는 대표적 플랫폼으로 자리 잡고 있어, 전통적 기사 중심 뉴스 소비와는 다른 방식으로 사건을 해석하고 반응하게 만든다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "2026년 3월 이후 이란 전쟁 관련 이슈는 한국 사회에서도 국제유가, LNG 공급, 중동 정세, 해상 안전과 같은 키워드와 함께 자주 보도되었다. 예컨대 “정부, 전쟁 충격 완화 위해 173억 달러 추경 추진”이라는 보도는 이란 전쟁 이슈가 한국의 경제와 정책 환경에도 직접 연결되는 국제 이슈로 인식되었음을 보여준다(곽민서·김아람, 연합뉴스, 2026년 3월 31일). 또한 한국에서 유튜브는 이미 주요 뉴스 소비 경로로 기능하고 있다. Digital News Report 2025를 인용한 보도에 따르면 한국의 유튜브 뉴스 이용률은 50% 수준으로 높게 나타났으며, 연령과 정치 성향에 따라 이용 양상에도 차이가 있는 것으로 보고되었다(김용석, 뉴스핌, 2025년 6월 17일).",
        indent=True,
    )
    add_paragraph(
        doc,
        "이러한 환경은 이란 전쟁과 같은 국제 사건이 한국 유튜브 뉴스 채널을 통해 어떤 관점으로 번역되고 재구성되는지, 그리고 그 결과 이용자들이 어떤 반응을 보이는지를 함께 살펴볼 필요성을 제기한다. 특히 유튜브 뉴스는 영상 자체뿐 아니라 제목, 설명, 채널 브랜드, 댓글이 하나의 해석 공간으로 결합되기 때문에, 동일한 국제 이슈가 채널에 따라 서로 다른 강조점을 띨 수 있다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "본 연구는 이란 전쟁 관련 한국 유튜브 뉴스 채널의 보도 프레임과 수용자 반응의 차이에 따른 플랫폼 재구성 양상을 측정하는 것을 목적으로 한다.",
        indent=True,
    )

    add_heading(doc, "II. 이론적 배경", 1)
    add_heading(doc, "1) 연구주제 현황", 2)
    add_paragraph(
        doc,
        "유튜브는 단순한 동영상 공유 플랫폼을 넘어 뉴스 유통과 소비, 반응의 형성이 동시에 이루어지는 플랫폼으로 발전해 왔다. 플랫폼 뉴스 환경에서는 기사 본문뿐 아니라 제목, 설명, 썸네일, 업로드 시점, 댓글 공간 등 다양한 요소가 함께 작동하며 이용자의 해석에 영향을 미친다. 특히 유튜브 뉴스는 이용자가 영상 시청 이전에 제목과 설명을 먼저 접하고, 시청 이후에는 댓글을 통해 다른 이용자의 반응과 감정 표현을 함께 확인한다는 특징이 있다. 따라서 유튜브 뉴스 연구에서는 기사 본문만이 아니라 플랫폼에서 실제로 노출되는 핵심 정보 단위를 함께 분석할 필요가 있다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "선행연구 또한 유튜브 뉴스 콘텐츠의 형식적 특성과 이용자 반응의 중요성을 지적해 왔다. 피연진·김경모(2023)는 지상파 방송사 유튜브 뉴스 콘텐츠가 전통 뉴스 기사와 다른 편집 양식을 지니며, 이용자 반응 역시 콘텐츠 포맷에 따라 달라진다고 보았다. 주은신(2020) 또한 지역 방송사 유튜브 뉴스 콘텐츠의 특성과 이용자 반응을 함께 분석하며, 플랫폼 기반 뉴스가 기존 방송 뉴스와는 다른 참여 구조를 형성한다고 설명하였다.",
        indent=True,
    )
    add_heading(doc, "2) 미디어 이론", 2)
    add_paragraph(
        doc,
        "프레이밍 이론은 언론이 사건의 특정 측면을 선택하고 강조함으로써 수용자가 현실을 해석하는 방식을 조직한다고 본다. Entman(1993)은 프레임을 문제 정의, 원인 진단, 도덕적 평가, 해결 방향 제시를 가능하게 하는 해석 틀로 설명하였다. 즉 프레임은 사건 자체보다 사건을 어떻게 보게 만드는지와 관련된다. de Vreese(2005)는 프레임을 범이슈적으로 적용 가능한 generic frame과 특정 사안에 맞게 구성되는 issue-specific frame으로 구분하였다. 이러한 구분은 본 연구가 기존 이론 범주를 단순 반복하기보다, 이란 전쟁이라는 특정 사안에 적합한 범주로 재구성할 수 있는 근거를 제공한다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "Semetko와 Valkenburg(2000)는 대표적 generic frame으로 갈등, 책임귀인, 경제적 결과, 인간흥미, 도덕성 프레임을 제시하였다. 본 연구는 이 다섯 범주를 이란 전쟁이라는 이슈에 맞게 조정하여 안보·군사, 국제정치·외교, 경제·에너지, 투자·시장, 인도주의·민간피해, 기타/혼합의 여섯 범주로 재구성한다. 안보·군사 프레임은 군사 공격과 보복, 확전, 안보 위협을 중심으로 나타나고, 국제정치·외교 프레임은 외교 협상, 국제법, 국제기구 대응, 국가 책임 논의를 중심으로 나타난다. 경제·에너지 및 투자·시장 프레임은 전쟁의 경제적 파급효과를 다루되, 전자는 유가·공급망·물가를, 후자는 증시·투자 심리·업종 반응을 중심으로 구분하였다. 인도주의·민간피해 프레임은 민간인 피해와 난민, 일상 파괴를 강조하는 경우에 적용하였다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "이처럼 프레이밍 이론은 국제 사건이 한국 유튜브 뉴스 안에서 어떤 방식으로 재구성되는지 설명하는 데 적절한 이론적 틀을 제공한다. 특히 유튜브 뉴스처럼 제목, 설명, 스크립트, 댓글이 함께 작동하는 플랫폼 환경에서는 프레임이 단순 기사 서술을 넘어 이용자 반응까지 연결되는 해석 구조로 기능할 가능성이 크다. 따라서 프레이밍 이론은 본 연구의 연구문제를 도출하고 보도 프레임과 수용자 반응의 관계를 해석하는 핵심 이론으로 활용될 수 있다.",
        indent=True,
    )

    add_heading(doc, "III. 연구문제", 1)
    add_paragraph(
        doc,
        "앞서 논의한 내용을 바탕으로, 본 연구는 2026년 3월 이후 이란 전쟁 관련 한국 유튜브 뉴스가 어떠한 프레임으로 사건을 구성하는지, 그리고 댓글 반응은 그러한 보도 프레임과 어떤 차이를 보이는지를 탐색하고자 한다. 본 연구는 국제 사건 자체의 진위를 판단하기보다, 유튜브 뉴스라는 플랫폼 안에서 사건이 어떻게 재구성되고 수용되는지에 초점을 둔다.",
        indent=True,
    )

    add_heading(doc, "연구문제 1. 이란 전쟁 관련 한국 유튜브 뉴스 채널의 보도 프레임과 주제는 어떠한가?", 2)
    add_paragraph(
        doc,
        "첫째, 동일한 국제 사건이라 하더라도 언론이 어떤 측면을 선택하고 강조하는지에 따라 수용자의 해석은 달라질 수 있다. 프레이밍 이론은 언론이 사건의 특정 측면을 선택하고 강조하는 방식을 통해 현실 인식을 조직한다고 본다. 따라서 한국 유튜브 뉴스 채널이 이란 전쟁 이슈를 어떤 프레임과 주제로 주로 재구성하는지 살펴보는 것은 본 연구의 가장 기초적인 질문이 된다.",
        indent=True,
    )

    add_heading(doc, "연구문제 2. 이란 전쟁 관련 유튜브 뉴스 영상의 댓글 반응 유형은 어떠한가?", 2)
    add_paragraph(
        doc,
        "둘째, 플랫폼 뉴스에서는 뉴스 생산물과 이용자 반응이 분리되지 않고 한 공간에서 함께 노출된다. 유튜브 댓글은 전체 여론을 대표하지는 않지만, 해당 뉴스 영상이 플랫폼 내에서 어떤 반응을 유발했는지를 보여주는 가시적 자료이다. 따라서 댓글 반응 유형의 분포를 분석함으로써 유튜브 뉴스의 수용 양상을 탐색할 필요가 있다.",
        indent=True,
    )

    add_heading(doc, "연구문제 3. 이란 전쟁 관련 한국 유튜브 뉴스 채널의 보도는 상대적으로 어떤 이념적 기울기를 보이는가?", 2)
    add_paragraph(
        doc,
        "셋째, 동일한 사건을 다루더라도 채널마다 강조하는 행위자, 위협의 성격, 해결 방향이 다를 수 있으며, 이는 텍스트 안에 상대적 이념적 기울기 차이로 나타날 수 있다. 본 연구에서 이념적 기울기는 채널의 본질적 정치 성향을 판정하는 라벨이 아니라, 이란 전쟁이라는 특정 이슈 맥락에서 드러난 상대적 해석 방향을 의미한다. 따라서 채널별 텍스트 안에서 진보적·보수적 단서가 어떤 비율로 나타나는지 비교할 필요가 있다.",
        indent=True,
    )

    add_heading(doc, "연구문제 4. 제목·설명 기반 주제와 실제 스크립트 기반 주제는 어떤 차이를 보이는가?", 2)
    add_paragraph(
        doc,
        "넷째, 유튜브 뉴스는 본문 내용 이전에 제목과 설명이 먼저 소비되기 때문에, 채널이 사건을 포장하는 방식과 실제 보도 내용이 다를 가능성이 있다. 플랫폼 뉴스 환경에서는 헤드라인 프레이밍과 본문 프레이밍의 차이가 이용자 해석에 중요한 영향을 미칠 수 있다. 따라서 제목·설명 텍스트와 스크립트 본문을 분리해 각각의 주제 구성을 비교하는 것은 유튜브 뉴스의 플랫폼적 특성을 드러내는 데 유용하다.",
        indent=True,
    )

    add_heading(doc, "연구모형", 2)
    add_bullets(
        doc,
        [
            "독립변인: 채널 유형, 보도 프레임, 제목·설명 기반 주제, 스크립트 기반 주제",
            "종속변인: 댓글 반응 유형, 상대적 이념적 기울기",
            "보조 변인: 게시 시기, 채널별 업로드량, 댓글 수",
        ],
    )

    add_heading(doc, "IV. 연구방법", 1)
    add_paragraph(
        doc,
        "본 연구는 이란 전쟁 관련 한국 유튜브 뉴스 채널의 보도 프레임과 수용자 반응의 차이에 따른 플랫폼 재구성 양상을 측정하는 것을 목적으로 한다.",
        indent=True,
    )
    add_paragraph(
        doc,
        f"분석대상은 {data['period_start']}부터 {data['period_end']}까지 업로드된 이란 전쟁 관련 한국어 유튜브 뉴스 영상이며, 최종 표본은 12개 채널의 영상 10,437개와 그에 결합된 최상위 댓글 233,691개로 구성하였다. 분석 대상 채널은 공영·지상파 방송사 계열, 종합편성채널 계열, 보도·경제전문 채널 계열로 구분하였고, 일반 시사 잡담 채널이나 개인 1인 미디어 채널은 제외하였다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "자료수집은 YouTube Data API를 활용하여 수행하였으며, 영상 제목·설명·게시일·채널명·조회수·댓글 수와 같은 메타데이터를 우선 확보하였다. 이후 공개 자막이나 확보 가능한 스크립트를 수집하였고, 스크립트가 불완전한 경우에는 별도의 텍스트 정제 규칙을 적용하여 분석 가능한 수준으로 보완하였다. 댓글 분석에서는 각 영상에 달린 최상위 댓글을 전수 수집하고, 광고성 댓글, 중복 댓글, URL-only 댓글, 15자 미만의 무의미한 댓글을 제외하였다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "분석 단위는 영상 1개와 댓글 1개이다. 영상 단위에서는 제목, 설명, 스크립트, 게시 시기, 채널명, 프레임 분류 결과, 주제 분류 결과, 이념적 기울기 점수를 종합하여 해석하였다. 댓글 단위에서는 댓글 텍스트를 기반으로 지지·응원, 정보보완·해설, 불안·공포, 비판·분노, 조롱·냉소, 기타·혼합의 여섯 범주로 반응 유형을 분류하였다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "보도 프레임은 기존 프레이밍 이론을 참고하되 이란 전쟁이라는 특정 이슈에 적합한 issue-specific frame으로 재구성하였다. 하나의 영상에는 가장 중심적으로 드러나는 프레임 하나를 우선 부여하였으며, 둘 이상의 프레임이 동시에 강하게 나타나거나 정보성 멘트와 채널 안내 문구가 과도하게 결합된 경우에는 기타/혼합 범주로 분류하였다. 이념적 기울기 점수는 진보적 단서와 보수적 단서의 상대적 출현 비율을 바탕으로 산출하였으며, 이 수치는 채널의 본질적 정치 성향이 아니라 이란 전쟁 보도 맥락에서 드러난 상대적 위치 추정치로 해석하였다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "주제 분석은 제목·설명 기반 텍스트와 스크립트 기반 텍스트를 분리하여 각각 수행하였다. 제목·설명 기반 주제는 채널이 영상을 어떤 방식으로 포장하고 소개하는지를 보여주는 지표로, 스크립트 기반 주제는 실제 보도 본문에서 어떤 사건과 표현이 반복되는지를 보여주는 지표로 해석하였다.",
        indent=True,
    )

    table1_rows = frame_table_rows()
    add_table(doc, ["프레임 유형", "내용", "출처"], table1_rows, "<표 1> 보도 프레임 유형 및 내용")

    add_heading(doc, "IV-1. 측정도구(코딩가이드) 초안", 1)
    add_paragraph(
        doc,
        "본 연구의 측정도구는 설문지가 아니라 내용분석용 코딩가이드이다. 코딩가이드는 영상 기본 정보, 보도 프레임, 주제 분류, 댓글 반응 분류 항목으로 구성하였다. 영상 코딩가이드에는 영상 번호, 채널명, 업로드 일자, 제목, URL, 채널 유형, 주요 프레임, 보조 메모가 포함되며, 댓글 코딩가이드에는 댓글 ID, 영상 ID, 댓글 작성일시, 최상위 댓글 여부, 제외 사유, 댓글 반응 유형 등이 포함된다. 자세한 항목은 본문 말미의 부록에 제시하였다.",
        indent=True,
    )

    add_heading(doc, "V. 연구결과", 1)
    add_heading(doc, "1) 사전조사(pilot study)", 2)
    add_paragraph(
        doc,
        "본코딩 이전에 검증용 표본을 추출하여 프레임 분류, 이념적 기울기 추정, 댓글 반응 분류 기준이 적절하게 작동하는지 점검하였다. 예비 검토 결과, Shorts·현장영상·플랫폼 안내 문구가 많은 영상은 자동 분류 과정에서 기타/혼합 범주로 과잉 분류되는 경향이 확인되었다. 이에 따라 불용어 사전과 텍스트 정제 규칙을 보완하였고, 제목·설명 기반 주제와 스크립트 기반 주제를 분리해 분석하도록 절차를 수정하였다. 본코딩에는 이러한 수정 규칙을 반영하였다.",
        indent=True,
    )

    add_heading(doc, "2) 코더 간 신뢰도 및 검증 절차", 2)
    add_paragraph(
        doc,
        "본 연구는 대규모 유튜브 코퍼스를 대상으로 한 규칙 기반·텍스트 기반 자동분류 체계를 중심으로 수행되었기 때문에, 전통적인 2인 수작업 코딩 방식의 홀스티 계수를 산출하지는 못하였다. 대신 프레임 분류, 이념적 기울기 추정, 댓글 반응 분류 각각에 대해 검증용 표본을 구성하여 규칙의 적절성과 경계 사례를 점검하는 방식의 예비 검증 절차를 수행하였다. 따라서 아래 표는 정식 상호신뢰도 계수표라기보다, 검증 절차와 보완 방향을 요약한 표로 해석할 필요가 있다. 최종 제출본에서는 가능하다면 별도의 수작업 2인 코딩을 통해 홀스티 계수를 추가 산출하는 것이 바람직하다.",
        indent=True,
    )
    add_table(
        doc,
        ["측정항목", "검증 방식", "비고"],
        [
            ["보도 프레임", "검증용 표본 수작업 점검", "Shorts·현장영상·플랫폼 문구 과잉 분류 여부 확인"],
            ["이념적 기울기", "검증용 표본 수작업 점검", "진보·보수 단서 사전의 방향성 적절성 검토"],
            ["댓글 반응", "검증용 표본 수작업 점검", "기타/혼합 과다 분류와 정보보완/해설 경계 사례 수정"],
        ],
        "<표 2> 검증 절차 요약",
    )

    add_heading(doc, "3) 표본의 특성", 2)
    type_rows = []
    for row in data["type_summary"].itertuples(index=False):
        type_rows.append([
            row.channel_type,
            str(row.channel_count),
            f"{int(row.video_count):,}",
            f"{int(row.comment_count):,}",
        ])
    add_table(
        doc,
        ["채널 유형", "채널 수", "영상 수", "댓글 수"],
        type_rows,
        "<표 3> 채널 유형별 표본 특성",
    )

    add_table(
        doc,
        ["채널명", "채널 유형", "영상 수", "대표 프레임", "대표 주제", "대표 댓글 반응"],
        [
            [
                row.channel_name,
                row.channel_type,
                f"{int(row.video_count):,}",
                row.dominant_frame,
                TOPIC_LABEL_MAP.get(str(row.dominant_topic), str(row.dominant_topic)),
                row.dominant_audience_reaction,
            ]
            for row in data["summary"].sort_values("video_count", ascending=False).itertuples(index=False)
        ],
        "<표 4> 채널별 핵심 특성",
    )

    add_heading(doc, "4) 연구문제별 결과", 2)

    add_heading(doc, "연구문제 1. 이란 전쟁 관련 한국 유튜브 뉴스 채널의 보도 프레임과 주제는 어떠한가?", 3)
    add_paragraph(
        doc,
        "전체 10,437개 영상 가운데 가장 많이 나타난 프레임은 안보·군사 프레임으로 4,692개(44.96%)였다. 그 다음은 기타/혼합 2,968개(28.44%), 경제·에너지 1,706개(16.35%), 투자·시장 661개(6.33%), 국제정치·외교 367개(3.52%), 인도주의·민간피해 43개(0.41%) 순으로 나타났다. 즉 한국 유튜브 뉴스 채널은 이란 전쟁을 전반적으로 군사 충돌과 안보 위협의 문제로 재구성하는 경향이 가장 강했다.",
        indent=True,
    )
    add_table(
        doc,
        ["프레임", "영상 수", "비율"],
        [
            [row.primary_frame, f"{int(row.video_count):,}", f"{row.share*100:.2f}%"]
            for row in data["frame_totals"].itertuples(index=False)
        ],
        "<표 5> 전체 표본의 프레임 분포",
    )
    add_paragraph(
        doc,
        "채널별로 보면 공영·지상파 계열과 종합편성채널 계열에서는 안보·군사 프레임이 대표 프레임으로 나타난 채널이 많았다. 반면 보도·경제전문 채널 계열에서는 YTN만 안보·군사 프레임이 대표적이었고, SBS Biz 뉴스·한국경제TV·매일경제TV는 투자·시장 프레임이, 연합뉴스TV는 기타/혼합 프레임이 대표적으로 나타났다. 이는 같은 이란 전쟁 이슈라도 채널 유형에 따라 전쟁 자체, 시장 충격, 혼합 보도 양상이 다르게 재구성될 수 있음을 보여준다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "제목·설명 기반 주제 분석에서는 이란·이스라엘 군사 충돌(topic_00), 트럼프와 미국 대응(topic_02), 호르무즈 해협과 원유 수송(topic_03), 전쟁의 시장·투자 충격(topic_04)이 상위 주제로 나타났다. 다만 채널 안내·부가 문구(topic_05), 시사 프로그램·플랫폼 문구(topic_06)처럼 유튜브 뉴스의 플랫폼적 포장 방식이 주제로 포착되는 경우도 확인되었다. 이는 유튜브 뉴스에서 제목과 설명이 사건의 본문 내용뿐 아니라 채널 브랜드와 편집 관행을 함께 드러낸다는 점을 보여준다.",
        indent=True,
    )
    add_table(
        doc,
        ["제목·설명 기반 주제", "문서 수", "비율", "대표 키워드"],
        [
            [
                row.topic_name,
                f"{int(row.document_count):,}",
                f"{row.share_of_documents*100:.2f}%",
                str(row.top_terms).split(",")[0:4].__str__().replace("[", "").replace("]", "").replace("'", ""),
            ]
            for row in data["topic"].head(5).itertuples(index=False)
        ],
        "<표 6> 제목·설명 기반 상위 주제",
    )

    add_heading(doc, "연구문제 2. 이란 전쟁 관련 유튜브 뉴스 영상의 댓글 반응 유형은 어떠한가?", 3)
    add_paragraph(
        doc,
        "전체 233,691개 댓글 가운데 가장 높은 비중을 차지한 반응은 기타/혼합 124,985개(53.48%)였다. 다음으로는 정보보완/해설 88,562개(37.90%), 비판/분노 12,384개(5.30%), 지지/응원 3,569개(1.53%), 불안/공포 3,383개(1.45%), 조롱/냉소 808개(0.35%) 순으로 나타났다. 즉 댓글 공간에서는 특정 입장에 대한 단순 지지보다, 정보 보완이나 해설형 반응, 혹은 명확히 분류하기 어려운 혼합 반응이 더 많이 나타났다.",
        indent=True,
    )
    add_table(
        doc,
        ["댓글 반응 유형", "댓글 수", "비율"],
        [
            [row.primary_reaction, f"{int(row.comment_count):,}", f"{row.share*100:.2f}%"]
            for row in data["reaction_totals"].itertuples(index=False)
        ],
        "<표 7> 전체 표본의 댓글 반응 분포",
    )
    add_paragraph(
        doc,
        "채널별로는 거의 모든 채널에서 기타/혼합과 정보보완/해설이 가장 큰 비중을 차지했지만, 세부 분포에는 차이가 있었다. MBCNEWS는 비판/분노 비율이 7.48%로 비교적 높았고, YTN과 JTBC News는 정보보완/해설 비중이 각각 39.35%, 39.15%로 높게 나타났다. 이는 플랫폼 내 댓글 반응이 일방적 지지나 반대보다 사건 해석과 보완 정보 제공 기능도 수행하고 있음을 시사한다.",
        indent=True,
    )

    add_heading(doc, "연구문제 3. 이란 전쟁 관련 한국 유튜브 뉴스 채널의 보도는 상대적으로 어떤 이념적 기울기를 보이는가?", 3)
    add_paragraph(
        doc,
        "이념적 기울기 추정 결과, 12개 채널 중 11개 채널이 이란 전쟁 이슈에서 보수적 기울기 범주에 속했고, SBS Biz 뉴스만 혼합/중간 범주로 분류되었다. 가장 높은 보수적 기울기 점수는 YTN(0.404)이었고, 그 다음은 연합뉴스TV(0.306), KBS News(0.244), 매일경제TV(0.235), SBS 뉴스(0.230) 순이었다. 다만 이 수치는 채널의 본질적 정치 성향을 판정하는 값이 아니라, 해당 이슈 보도에서 드러난 상대적 해석 방향이라는 점에서 제한적으로 해석해야 한다.",
        indent=True,
    )
    add_table(
        doc,
        ["채널명", "진보 단서 수", "보수 단서 수", "상대 점수", "해석"],
        [
            [
                row.channel_name,
                f"{int(row.progressive_cue_hits):,}",
                f"{int(row.conservative_cue_hits):,}",
                f"{row.ideology_relative_score:.3f}",
                row.ideology_relative_label,
            ]
            for row in data["ideology_sorted"].head(12).itertuples(index=False)
        ],
        "<표 8> 채널별 이념적 기울기 추정 결과",
    )
    add_paragraph(
        doc,
        "채널별 이념적 기울기 차이는 프레임 분포와도 일정 부분 연결되었다. 예컨대 시장 중심 보도가 많은 경제전문 채널은 상대적으로 이념 단서가 약하게 나타났고, 군사·안보 중심 보도가 강한 채널에서는 보수적 기울기 점수가 더 높게 산출되는 경향이 있었다. 이는 프레임과 이념적 기울기를 동일 개념으로 보아서는 안 되지만, 특정 프레임 배치가 상대적 해석 방향과 결합할 수 있음을 보여준다.",
        indent=True,
    )

    add_heading(doc, "연구문제 4. 제목·설명 기반 주제와 실제 스크립트 기반 주제는 어떤 차이를 보이는가?", 3)
    add_paragraph(
        doc,
        "스크립트 기반 주제 분석에서는 호르무즈 해협과 원유 수송(topic_04), 트럼프와 미국 대응(topic_03), 미사일·드론 공격(topic_00), 최고지도자와 내부 권력(topic_07)이 상위 주제로 나타났다. 이는 제목·설명 기반 주제가 채널 안내·플랫폼 문구와 같은 포장 효과를 일부 포함하는 것과 달리, 스크립트 기반 주제는 실제 본문 내용의 반복 이슈를 더 직접적으로 반영한다는 점을 보여준다. 동시에 두 분석 모두 이란·이스라엘 군사 충돌과 호르무즈 해협, 트럼프 대응을 핵심 축으로 포착했다는 점에서, 메타데이터와 본문 주제가 완전히 분리되기보다는 핵심 이슈를 공유하면서도 강조 지점이 달라진다고 볼 수 있다.",
        indent=True,
    )
    add_table(
        doc,
        ["스크립트 기반 주제", "문서 수", "비율", "대표 키워드"],
        [
            [
                row.topic_name,
                f"{int(row.document_count):,}",
                f"{row.share_of_documents*100:.2f}%",
                str(row.top_terms).split(",")[0:4].__str__().replace("[", "").replace("]", "").replace("'", ""),
            ]
            for row in data["topic_script"].head(5).itertuples(index=False)
        ],
        "<표 9> 스크립트 기반 상위 주제",
    )

    add_heading(doc, "VI. 결론 및 함의", 1)
    add_paragraph(
        doc,
        "첫째, 본 연구는 2026년 3월 이후 이란 전쟁 관련 한국 유튜브 뉴스 채널 보도가 전반적으로 안보·군사 프레임을 가장 많이 활용하며, 댓글 반응에서는 기타/혼합과 정보보완/해설 반응이 중심을 이룬다는 점을 확인하였다. 또한 채널별 분석 결과, 공영·지상파 및 종합편성채널은 군사·안보 중심 보도가 강했고, 일부 경제전문 채널은 투자·시장 프레임을 상대적으로 더 많이 활용하였다. 이념적 기울기 추정 결과에서는 대부분의 채널이 이란 전쟁 이슈에서 상대적으로 보수적 방향에 가까운 텍스트 단서를 보였으며, 이는 프레임 배치와 완전히 동일하지는 않지만 일정한 연결 가능성을 보여주었다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "둘째, 본 연구의 한계는 연구방법 측면에서 분명하다. 첫째, 본 연구는 규칙 기반·텍스트 기반 자동분류를 중심으로 수행되었기 때문에 전통적 의미의 2인 코더 간 신뢰도 수치를 충분히 제시하지 못하였다. 둘째, 댓글은 플랫폼 내에서 가시성이 높은 반응 자료이지만 전체 여론을 대표하지 않으며, 익명성과 감정 표현의 강도, 알고리즘 노출 효과를 동시에 포함한다. 셋째, 메타데이터와 스크립트 활용 범위를 넓혔음에도 불구하고 여전히 플랫폼 안내 문구, Shorts, 현장영상처럼 분석에 불리한 형식이 남아 있어 일부 주제와 프레임이 기타/혼합으로 과대 분류되었을 가능성이 있다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "셋째, 후속 연구에서는 두 가지 확장이 가능하다. 하나는 수작업 2인 코딩을 병행하여 프레임과 댓글 반응 분류의 신뢰도를 보다 엄밀하게 제시하는 것이다. 다른 하나는 제목·설명·썸네일·스크립트·댓글을 통합한 다층적 플랫폼 프레이밍 연구로 확장하여, 동일한 국제 사건이 한국 뉴스 플랫폼 안에서 어떻게 다층적으로 재가공되는지를 더 입체적으로 분석하는 것이다. 특히 알고리즘 추천 구조와 조회수, 업로드 시간대를 함께 고려하면 유튜브 뉴스의 노출 메커니즘과 수용자 반응의 관계를 더 정교하게 설명할 수 있을 것이다.",
        indent=True,
    )
    add_paragraph(
        doc,
        "마지막으로 이번 학기 수업을 돌아보면, 처음에는 뉴스 프레임 분석을 유튜브 데이터와 실제로 연결하는 것이 막막하게 느껴졌지만, 수집과 정제, 분석 과정을 직접 반복하면서 미디어 연구가 단지 이론 설명이 아니라 데이터를 통해 확인되는 과정이라는 점을 체감할 수 있었다. 특히 프레이밍 이론을 실제 유튜브 뉴스 제목과 스크립트, 댓글에 적용해 보면서 이론과 실제 텍스트 사이의 거리를 구체적으로 이해할 수 있었던 점이 가장 큰 보람이었다. 동시에 수작업 코딩과 자동분류의 차이, 데이터 품질과 플랫폼 구조의 영향, 시각화와 문서화의 중요성을 끝까지 체감했다는 점에서 아쉬움과 배움이 함께 남았다. 결과적으로 이번 프로젝트는 연구 질문을 데이터 구조와 연결하는 경험을 제공했다는 점에서 의미가 있었고, 다음에는 더 엄밀한 신뢰도 설계와 시계열 비교를 추가해 보고 싶다. 수업 전체로 보았을 때, 실제 연구의 전 과정을 경험했다는 점이 가장 큰 성과였다.",
        indent=True,
    )

    add_heading(doc, "VII. 참고문헌", 1)
    refs = [
        "곽민서·김아람. (2026년 3월 31일). 정부, 전쟁 충격 완화 위해 173억 달러 추경 추진. 연합뉴스.",
        "김용석. (2025년 6월 17일). 한국인의 유튜브 뉴스 이용률 50%…연령·성향별 차이도 뚜렷. 뉴스핌.",
        "주은신. (2020). 지역 방송사 유튜브 뉴스 콘텐츠 특성과 이용자 반응에 관한 내용분석. 한국콘텐츠학회논문지, 20(9), 169-186. https://doi.org/10.5392/JKCA.2020.20.09.169",
        "피연진·김경모. (2023). 지상파 방송사 유튜브 영상 뉴스의 특성과 이용 반응: KBS 크랩, MBC 14F, SBS 스브스뉴스의 내용분석. 한국방송학보, 37(4), 334-372.",
        "de Vreese, C. H. (2005). News framing: Theory and typology. Information Design Journal + Document Design, 13(1), 51-62. https://doi.org/10.1075/idjdd.13.1.06vre",
        "Entman, R. M. (1993). Framing: Toward clarification of a fractured paradigm. Journal of Communication, 43(4), 51-58. https://doi.org/10.1111/j.1460-2466.1993.tb01304.x",
        "Reuters Institute for the Study of Journalism. (2025). Digital News Report 2025. https://reutersinstitute.politics.ox.ac.uk/digital-news-report/2025",
        "Semetko, H. A., & Valkenburg, P. M. (2000). Framing European politics: A content analysis of press and television news. Journal of Communication, 50(2), 93-109. https://doi.org/10.1111/j.1460-2466.2000.tb02843.x",
    ]
    for ref in refs:
        add_paragraph(doc, ref)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "부록 1. 영상 코딩가이드", 1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["영상 기본 정보", "영상번호, 채널명, 업로드일자, 제목, video_id, 채널 유형"],
            ["보도 프레임", "안보·군사, 국제정치·외교, 경제·에너지, 투자·시장, 인도주의·민간피해, 기타/혼합"],
            ["주제 분석", "제목·설명 기반 주제, 스크립트 기반 주제를 별도로 분류"],
            ["보조 메모", "복합 프레임이거나 분류가 어려운 경우 메모 기록"],
        ],
        "<부록표 1> 영상 코딩가이드 요약",
    )

    add_heading(doc, "부록 2. 댓글 코딩가이드", 1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["댓글 기본 정보", "댓글 ID, 영상번호, 작성일시, 최상위 댓글 여부"],
            ["제외 기준", "광고성, 무의미한 반복, URL-only, 분석 불가 댓글"],
            ["반응 유형", "지지·응원, 정보보완·해설, 불안·공포, 비판·분노, 조롱·냉소, 기타/혼합"],
            ["비고", "중복 작성자 여부와 경계 사례 메모"],
        ],
        "<부록표 2> 댓글 코딩가이드 요약",
    )

    add_heading(doc, "부록 3. 제출 전 확인 필요 사항", 1)
    add_bullets(
        doc,
        [
            "서론의 언론보도 인용 2건은 기자명과 날짜를 최종 한 번 더 확인할 것.",
            "코더 간 신뢰도는 현재 검증 절차 중심으로 서술되어 있으므로, 가능하다면 수작업 2인 코딩 후 홀스티 계수를 추가할 것.",
            "지도교수의 모델논문 요구에 맞춰 모델논문 PDF 1편을 별도 첨부할 것.",
            "피드백된 1페이지 초안, 피드백된 중간보고서, 수집 데이터 파일(SPSS 대체 가능 형식 포함)을 함께 제출할 것.",
        ],
    )

    return doc


def main() -> None:
    doc = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
