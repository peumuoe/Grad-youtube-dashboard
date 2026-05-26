from __future__ import annotations

import re
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE_MD = ROOT / "data" / "processed" / "final_report_polished_draft.md"
DEFAULT_OUTPUT_DIR = ROOT / "outputs" / "report"
DEFAULT_OUTPUT_DOCX = DEFAULT_OUTPUT_DIR / "final_report_polished_draft.docx"


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(10.5)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10.5)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10.5)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
    elif level == 2:
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
    else:
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def export_md_to_docx(source_md: Path, output_docx: Path) -> None:
    text = source_md.read_text(encoding="utf-8")
    lines = text.splitlines()

    document = Document()
    set_default_font(document)
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    title_done = False

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            document.add_paragraph()
            continue
        if line == "---":
            continue

        if line.startswith("# "):
            if not title_done:
                title_done = True
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(line[2:].strip())
                run.bold = True
                run.font.size = Pt(18)
                run.font.name = "Malgun Gothic"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            else:
                add_heading(document, line[2:].strip(), 1)
            continue

        if line.startswith("## "):
            add_heading(document, line[3:].strip(), 1)
            continue

        if line.startswith("### "):
            add_heading(document, line[4:].strip(), 2)
            continue

        if line.startswith("- "):
            add_bullet(document, line[2:].strip())
            continue

        # Strip simple markdown emphasis
        cleaned = re.sub(r"\*([^*]+)\*", r"\1", line)
        add_paragraph(document, cleaned)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE_MD)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT_DOCX)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    source = args.source if args.source.is_absolute() else ROOT / args.source
    output = args.output if args.output.is_absolute() else ROOT / args.output
    export_md_to_docx(source, output)
    print(f"Saved: {output}")


if __name__ == "__main__":
    main()
